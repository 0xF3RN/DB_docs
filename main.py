import psycopg2
import configparser
from rich import print
import tabulate
from art import text2art
import os
import docx


#подключение к бд
try:
    config = configparser.ConfigParser()
    config.read("config.cfg")
    db_params = config["PostgreSQL"]
    connection = psycopg2.connect(host=db_params["host"], dbname=db_params["name"],
                                user=db_params["user"], password=db_params["password"],
                                port=db_params["port"])
except Exception:
    print ("[ОШИБКА] Невозможно подключиться к БД, проверьте данные в config.cfg\n Иначе проверьте состояние БД")
    

def clear_screen():
    """Функция для отчистки экрана"""
    if os.name == 'nt':
        os.system('cls')
    else:
        os.system('clear')


def get_data_for_invoice():
    """Эта функция выводит список машин, которые были арендованы и кем, для последующей передчи их в выходной документ"""
    clear_screen()
    cur = connection.cursor()
    cur.execute("SET lc_monetary TO 'ru_RU.UTF-8'")
    cur.execute('''
SELECT rent_id, manufacturer, CONCAT(client_surname , ' ',client_name  , ' ', client_fathersname) AS fullname,
start_of_rent, end_of_rent, rent_cost FROM client join rent on id_client = client_id join invoice on
rent_id = id_rent join automobile on id_automobile = automobile_id;
''')
    data = cur.fetchall()
    columns = ["ID", "Car Model", "Client", "Date start", "Date end", "Cost"]
    table = tabulate.tabulate(data, columns, tablefmt="fancy")
    print(table,'\n----  -----------  ------------------------  ------------  ----------  ------------')
    print("\n Выбирите ID, человека для которого вы хотите сделать выходной документ: ", end='')
    x = int(input())
    try:
        cur.execute(f'''
SELECT rent_id, manufacturer, CONCAT(client_surname , ' ',client_name  , ' ', client_fathersname) AS fullname,
start_of_rent, end_of_rent, rent_cost FROM client join rent on id_client = client_id join invoice on
rent_id = id_rent join automobile on id_automobile = automobile_id WHERE rent_id={x};
''')    
        data = [cur.fetchone()]   
        table = tabulate.tabulate(data, columns, tablefmt="fancy")
        print(table,'\n----  -----------  ------------------------  ------------  ----------  ------------')
        choices = {
            1: lambda: create_rent_doc(data),
            2: lambda: menu()
        }
        x = int(input("Вы уверенны, что хотите сделать выходной документ для этого человека?\n 1 - 'Да'; 2 - 'Нет': "))
        choices[x]()
    except Exception as e:
        print(" \n[ОШИБКА] Ошибка ввода индекса", e)
    connection.commit()
    cur.close()
    connection.close()


def create_rent_doc(data:list):
    """Создание выходного документа: счет"""
    doc = docx.Document()
    doc.add_heading('Счет за авто', 0)
    for row in data:
        doc.add_paragraph(f'Арендованная машина: {row[1]}')
        doc.add_paragraph(f'Клиент: {row[2]}')
        doc.add_paragraph(f'Дата начала аренды: {row[3].strftime("%d.%m.%Y")}')
        doc.add_paragraph(f'Дата конца аренды: {row[4].strftime("%d.%m.%Y")}')
        doc.add_paragraph(f'К оплате: {row[5]}')
        doc.add_paragraph(f'Подпись клиента:_________________')
    doc.add_heading('', 0)
    doc.save('Выходные документы\\счет клиенту.docx')
    print(27*'-')
    print(text2art("DONE"), 27*'-')
    

def menu():
    """Главное меню"""
    clear_screen()
    choices = {
        1:lambda:get_data_for_invoice(),
        2:lambda:get_data_for_claim(),
        3:lambda:get_data_for_service(),
        4:lambda:print(50*'-', '\n',text2art("Goodbye"), 50*'-'),
        5:lambda:print(50*'-', '\n',text2art("Goodbye"), 50*'-')
    }
    print('\n',62*'-', '\n',text2art("Documents"), 62*'-')

    print("1. Счет клиенту\n2. Претензия к клиенту\n3. Направление на ТО\n4. Выход из программы")
    x = int(input("Какой документ вы хотите получить(выберите цифру): "))
    choices[x]()
    

def get_data_for_service():
    """Мне лень писать тут доку"""
    clear_screen()
    cur = connection.cursor()
    cur.execute('''
        SELECT  id_automobile, manufacturer, organization_name, legal_address, type_of_work, day_of_work
        FROM automobile
        JOIN service ON id_automobile = automobile_id
        JOIN type_of_work ON type_of_work_id = id_type_of_work
        join partner on parter_id = id_partner
        WHERE now()::date - day_of_work::date > 365;

    ''')
    columns = ["ID", "Car Model", "Organization", "Address", "Type of work", "Last work"]
    table = tabulate.tabulate(cur.fetchall(), columns, tablefmt="fancy")
    print(table,'\n----  ---------------------  --------------  ---------  --------------  -----------')
    print("\n Выбирите ID, человека для которого вы хотите сделать выходной документ: ", end='')
    x = int(input())
    try:
        cur.execute(f'''
        SELECT  id_automobile, manufacturer, organization_name, legal_address, type_of_work, day_of_work
        FROM automobile
        JOIN service ON id_automobile = automobile_id
        JOIN type_of_work ON type_of_work_id = id_type_of_work
        join partner on parter_id = id_partner
        WHERE now()::date - day_of_work::date > 365 and id_automobile = {x} ;
''')    
        data = [cur.fetchone()]   
        table = tabulate.tabulate(data, columns, tablefmt="fancy")
        print(table,'\n----  ---------------------  --------------  ---------  --------------  -----------')
        choices = {
            1: lambda: create_service_doc(data),
            2: lambda: menu()
        }
        x = int(input("Вы уверенны, что хотите сделать выходной документ для этой машины?\n 1 - 'Да'; 2 - 'Нет': "))
        choices[x]()
    except Exception:
        print(" \n[ОШИБКА] Ошибка ввода индекса")
    connection.commit()
    cur.close()
    connection.close()


def create_service_doc(data:list):
    """Создание выходного документа: ТО"""
    doc = docx.Document()
    doc.add_heading('Техническое обслуживание', 0)
    for row in data:
        doc.add_paragraph(f'Машина: {row[1]}')
        doc.add_paragraph(f'Организация партнер: {row[2]}')
        doc.add_paragraph(f'Адресс: {row[3]}')
        doc.add_paragraph(f'Вид работы: {row[4]}')
        doc.add_paragraph(f'Дата последней работы: {row[5].strftime("%d.%m.%Y")}')
        doc.add_paragraph(f'Подпись отвественного лица:_________________')
    doc.add_heading('', 0)
    doc.save('Выходные документы\\ТО.docx')
    print(27*'-')
    print(text2art("DONE"), 27*'-')


def get_data_for_claim():
    """и тут мне лень писать тут доку тож"""
    clear_screen()
    cur = connection.cursor()
    cur.execute('''
    SELECT id_claim, manufacturer, CONCAT(client_surname , ' ',client_name  , ' ', client_fathersname)
    AS fullname, date_of_creation, description  FROM client join rent on id_client = client_id join
    automobile on id_automobile = automobile_id join claim on id_rent = rent_id;
    ''')
    columns = ["ID", "Car Model", "Client", "Date of creation", "Description"]
    table = tabulate.tabulate(cur.fetchall(), columns, tablefmt="fancy")
    print(table,'\n----  -----------  ------------------------  ------------------  -------------')
    print("\n Выбирите ID, человека для которого вы хотите сделать выходной документ: ", end='')
    x = int(input())
    try:
        cur.execute(f'''
    SELECT id_claim, manufacturer, CONCAT(client_surname , ' ',client_name  , ' ', client_fathersname)
    AS fullname, date_of_creation, description  FROM client join rent on id_client = client_id join
    automobile on id_automobile = automobile_id join claim on id_rent = rent_id WHERE id_claim = {x};
''')    
        data = [cur.fetchone()]   
        table = tabulate.tabulate(data, columns, tablefmt="fancy")
        print(table,'\n----  -----------  ------------------------  ------------------  -------------')
        choices = {
            1: lambda: create_claim_doc(data),
            2: lambda: menu()
        }
        x = int(input("Вы уверенны, что хотите сделать выходной документ для этого человека?\n 1 - 'Да'; 2 - 'Нет': "))
        choices[x]()
    except Exception:
        print(" \n[ОШИБКА] Ошибка ввода индекса")
    connection.commit()
    cur.close()
    connection.close()

 
def create_claim_doc(data:list):
    """Создание выходного документа: претензия"""
    doc = docx.Document()
    doc.add_heading('Претензия клиенту', 0)
    for row in data:
        doc.add_paragraph(f'Арендованная машина: {row[1]}')
        doc.add_paragraph(f'Клиент: {row[2]}')
        doc.add_paragraph(f'Дата жалобы: {row[3].strftime("%d.%m.%Y")}')
        doc.add_paragraph(f'Причина жалобы: {row[4]}')
        doc.add_paragraph(f'Подпись клиента:_________________')
    doc.add_heading('', 0)
    doc.save('Выходные документы\\претензия клиенту.docx')
    print(27*'-')
    print(text2art("DONE"), 27*'-')


def main():
    menu()


if __name__ == "__main__":
    main()